"""DOCX extractor for IKMAS ingestion system."""

from typing import List
from langchain_core.documents import Document
from app.rag.storage import StoredFile
from docx import Document as DocxDocument


def extract_docx_documents(stored: StoredFile) -> List[Document]:
    """
    Extract documents from a DOCX file.
    
    Args:
        stored: StoredFile object containing the DOCX file
        
    Returns:
        List of Document objects with metadata
    """
    # Read the DOCX file content
    doc = DocxDocument(stored.path)
    
    # Extract text from paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = '\n'.join(paragraphs)
    
    # Create a single document with metadata
    doc = Document(
        page_content=content,
        metadata={
            "file_id": stored.file_id,
            "source": stored.original_name,
            "chunk_type": "docx_text"
        }
    )
    
    return [doc]