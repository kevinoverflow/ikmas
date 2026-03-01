from typing import List, Tuple, Literal
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader

from app.rag.tokenizer import get_tokenizer
from app.rag.storage import save_upload, StoredFile

ConflictAction = Literal["skip", "replace", "rename"]


def uploads_to_bytes(files) -> List[Tuple[str, bytes]]:
    return [(f.name, f.getvalue()) for f in files]


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _new_splitter(chunk_size: int, chunk_overlap: int):
    tokenizer = get_tokenizer()
    return RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
        tokenizer,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        add_start_index=True,
    )


def _base_metadata(stored: StoredFile, project: str, artefact_type: str, seci_phase: str | None = None) -> dict:
    meta = {
        "file_id": stored.file_id,
        "source": stored.original_name,
        "project": project,
        "artefact_type": artefact_type,
        "timestamp": _now_iso(),
    }
    if seci_phase:
        meta["SECI_phase"] = seci_phase
    return meta


def _split_pdf(stored: StoredFile, splitter, base_meta: dict) -> List[Document]:
    docs = PyPDFLoader(str(stored.path)).load()
    for d in docs:
        d.metadata.update(base_meta)
    return splitter.split_documents(docs)


def _split_text(stored: StoredFile, splitter, base_meta: dict) -> List[Document]:
    content = Path(stored.path).read_text(encoding="utf-8", errors="ignore")
    doc = Document(page_content=content, metadata=base_meta)
    return splitter.split_documents([doc])


def _split_docx(stored: StoredFile, splitter, base_meta: dict) -> List[Document]:
    if DocxDocument is None:
        raise RuntimeError("python-docx is required for DOCX ingestion")
    docx = DocxDocument(str(stored.path))
    text = "\n".join(p.text for p in docx.paragraphs if p.text)
    doc = Document(page_content=text, metadata=base_meta)
    return splitter.split_documents([doc])


def split_file(
    stored: StoredFile,
    chunk_size: int = 512,
    chunk_overlap: int = 80,
    project: str = "default",
    artefact_type: str = "source_document",
    seci_phase: str | None = None,
):
    splitter = _new_splitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    suffix = stored.path.suffix.lower()
    base_meta = _base_metadata(stored, project=project, artefact_type=artefact_type, seci_phase=seci_phase)

    if suffix == ".pdf":
        return _split_pdf(stored, splitter, base_meta)
    if suffix == ".docx":
        return _split_docx(stored, splitter, base_meta)
    if suffix in {".md", ".txt"}:
        return _split_text(stored, splitter, base_meta)

    raise ValueError(f"Unsupported file type: {suffix}")


def split_pdf_file(stored: StoredFile, chunk_size=512, chunk_overlap=80):
    # Backward-compatible wrapper used by existing tests/UI.
    return split_file(stored, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def ingest_uploads(
    collection_id: str,
    uploaded_files,
    on_name_conflict: ConflictAction = "skip",
    chunk_size: int = 512,
    chunk_overlap: int = 80,
    project: str = "default",
    artefact_type: str = "source_document",
    seci_phase: str | None = None,
) -> Tuple[List, dict]:
    """
    Takes Streamlit UploadedFile objects, persists them, returns:
      - all_chunks: list of split LangChain Documents
      - stats: counts of saved/skipped/replaced/renamed/errors
    """
    all_chunks = []
    stats = {
        "saved": 0,
        "replaced": 0,
        "renamed": 0,
        "skipped_identical": 0,
        "skipped_conflict": 0,
        "errors": 0,
        "error_messages": [],
    }

    for f in uploaded_files:
        try:
            data = f.getvalue()
            status, stored = save_upload(
                collection_id=collection_id,
                filename=f.name,
                data=data,
                on_name_conflict=on_name_conflict,
            )

            if status in stats:
                stats[status] += 1
            else:
                stats["saved"] += 1

            if stored is None:
                continue

            chunks = split_pdf_file(
                stored,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
            for chunk in chunks:
                if not hasattr(chunk, "metadata") or not isinstance(chunk.metadata, dict):
                    continue
                chunk.metadata.setdefault("project", project)
                chunk.metadata.setdefault("artefact_type", artefact_type)
                chunk.metadata.setdefault("timestamp", _now_iso())
                if seci_phase:
                    chunk.metadata.setdefault("SECI_phase", seci_phase)
            all_chunks.extend(chunks)

        except Exception as e:
            stats["errors"] += 1
            stats["error_messages"].append(f"{getattr(f, 'name', '<unknown>')}: {e}")
            continue

    return all_chunks, stats
